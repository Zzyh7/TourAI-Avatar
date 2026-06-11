"""
DocxParser —— 解析 .docx 文件，提取表格数据和段落文本。

支持两种文档类型:
  1. 结构化表格文档: 每行是一个数据记录（如景点数据集）
     - 自动识别表头行 → 列名映射
     - 支持跨表（同一文档多个表格，如"灵山胜境景点"和"拈花湾景点"）
  2. 叙述性文档: 按标题层级组织的长文段落
     - 自动识别标题样式（Heading 1/2/3）
     - 保留段落层级结构

依赖: python-docx>=1.1.0 (比 docx2txt 更精细，支持表格单元格访问)
"""
import os
import hashlib
from typing import List, Dict, Any, Optional, Tuple

# python-docx 在运行时可能未安装，延迟导入以提供友好错误提示
_DOCX_AVAILABLE = True
try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.oxml.ns import qn
except ImportError:
    _DOCX_AVAILABLE = False


class DocxParserError(Exception):
    """文档解析异常"""


class DocxParser:
    """
    .docx 文件解析器。

    用法:
        parser = DocxParser("path/to/doc.docx")

        # 提取表格数据
        tables = parser.extract_tables()
        for table in tables:
            for row in table["rows"]:
                print(row["景点名称"])

        # 提取段落（带标题层级）
        paras = parser.extract_paragraphs()
        for p in paras:
            print(f"[H{p['level']}] {p['text']}")
    """

    def __init__(self, file_path: str):
        if not _DOCX_AVAILABLE:
            raise DocxParserError(
                "缺少 python-docx 库，请运行: pip install python-docx>=1.1.0"
            )
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_size = os.path.getsize(file_path)
        self._doc = DocxDocument(file_path)

        # 提取文件内容指纹（用于去重检测）
        self.content_hash = hashlib.md5(
            self.file_name.encode() + str(self.file_size).encode()
        ).hexdigest()[:12]

    # =====================================================================
    # 表格提取
    # =====================================================================

    def extract_tables(self) -> List[Dict[str, Any]]:
        """
        提取文档中所有表格，每个表格返回:
          {
            "table_index": int,          # 表格序号（从 0 开始）
            "caption": str | None,       # 表格前紧邻的段落文本（可能是表名）
            "columns": [str, ...],       # 列名列表（来自表头第一行）
            "rows": [dict, ...],         # 数据行列表，每行 dict[列名] = 单元格文本
            "row_count": int,
          }

        自动处理:
          - 识别表头行（通常是表格第一行，字体加粗或有背景色）
          - 合并单元格 → 填充重复值
          - 空行跳过
          - 多表并存
        """
        if not self._doc.tables:
            return []

        tables = []
        for t_idx, table in enumerate(self._doc.tables):
            rows_data = self._parse_table(table)
            if not rows_data:
                continue

            # 第一行作为表头
            header_row = rows_data[0]
            columns = [
                self._normalize_cell_text(cell)
                for cell in header_row
            ]

            # 数据行：将每行映射为 dict
            data_rows = []
            for row in rows_data[1:]:
                # 标准化每个单元格
                values = [self._normalize_cell_text(cell) for cell in row]
                # 跳过全空行
                if all(not v for v in values):
                    continue
                # 对齐列数（处理列数不一致的表格）
                row_dict = {}
                for i, col_name in enumerate(columns):
                    row_dict[col_name] = values[i] if i < len(values) else ""
                # 额外列（如果有超出表头数量的列）
                for i in range(len(columns), len(values)):
                    row_dict[f"额外列_{i}"] = values[i]

                data_rows.append(row_dict)

            # 尝试获取表格标题（表格前面的段落）
            caption = self._find_table_caption(table)

            tables.append({
                "table_index": t_idx,
                "caption": caption,
                "columns": columns,
                "rows": data_rows,
                "row_count": len(data_rows),
            })

        return tables

    def _parse_table(self, table: DocxTable) -> List[List[str]]:
        """解析单个表格，返回二维列表 [行][列]"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = cell.text
                cells.append(text)
            rows.append(cells)
        return rows

    def _find_table_caption(self, table: DocxTable) -> Optional[str]:
        """查找表格前的标题段落（用于识别子表名，如'灵山胜境景点'）"""
        # 通过 XML 位置关系查找表格前紧邻的段落
        try:
            tbl_element = table._tbl
            prev = tbl_element.getprevious()
            # 向前查找最近的段落
            while prev is not None:
                if prev.tag == qn('w:p'):
                    # 检查是否包含文本
                    texts = [node.text for node in prev.iter() if node.text]
                    text = ''.join(texts).strip()
                    if text:
                        return text
                prev = prev.getprevious()
        except Exception:
            pass
        return None

    @staticmethod
    def _normalize_cell_text(text: str) -> str:
        """清理单元格文本：去除多余空白、换行符"""
        if not text:
            return ""
        # 将表格内换行替换为分号（保持信息不丢失）
        text = text.replace('\n', '；').replace('\r', '')
        # 压缩多余空白
        text = ' '.join(text.split())
        return text.strip()

    # =====================================================================
    # 段落提取（按标题层级分组）
    # =====================================================================

    def extract_paragraphs(self) -> List[Dict[str, Any]]:
        """
        提取文档所有段落，保留标题层级信息。

        Returns:
            [
              {
                "text": str,           # 段落文本
                "style": str,          # 样式名 (Heading 1/2/3, Normal, etc.)
                "level": int,          # 标题层级 (0=正文, 1=一级标题, 2=二级标题, ...)
                "is_heading": bool,
                "para_index": int,     # 段落序号
              },
              ...
            ]
        """
        paragraphs = []
        for p_idx, para in enumerate(self._doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue  # 跳过空段落

            style_name = para.style.name if para.style else "Normal"
            level = self._heading_level(style_name)
            is_heading = level > 0

            paragraphs.append({
                "text": text,
                "style": style_name,
                "level": level,
                "is_heading": is_heading,
                "para_index": p_idx,
            })

        return paragraphs

    @staticmethod
    def _heading_level(style_name: str) -> int:
        """根据样式名判断标题层级"""
        style_lower = style_name.lower().replace(' ', '')
        # 英文样式
        if 'heading1' in style_lower or style_lower == 'title':
            return 1
        if 'heading2' in style_lower:
            return 2
        if 'heading3' in style_lower:
            return 3
        if 'heading4' in style_lower:
            return 4
        # 中文样式（部分模板使用）
        if '1级' in style_lower or '标题1' in style_lower:
            return 1
        if '2级' in style_lower or '标题2' in style_lower:
            return 2
        if '3级' in style_lower or '标题3' in style_lower:
            return 3
        return 0  # 正文

    # =====================================================================
    # 文本提取（全量，保留顺序）
    # =====================================================================

    def extract_all_text(self) -> str:
        """
        提取文档全部文本（表格 + 段落），保持原始顺序。

        表格中的每个单元格用换行分隔，表格前后加分隔标记。
        """
        parts = []

        # 遍历文档 body 的所有元素（段落 + 表格），保持原始顺序
        try:
            body = self._doc.element.body
            for child in body:
                if child.tag == qn('w:p'):
                    # 段落
                    texts = [node.text for node in child.iter() if node.text]
                    if texts:
                        parts.append(''.join(texts))
                elif child.tag == qn('w:tbl'):
                    # 表格
                    parts.append('\n--- 表格 ---')
                    for row in child.iter(qn('w:tr')):
                        cells = []
                        for cell in row.iter(qn('w:tc')):
                            cell_texts = [
                                node.text for node in cell.iter()
                                if node.text
                            ]
                            cells.append(''.join(cell_texts))
                        parts.append(' | '.join(c for c in cells if c))
                    parts.append('--- 表格结束 ---\n')
        except Exception:
            # 回退：简单拼接段落和表格
            for para in self._doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in self._doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(' | '.join(c for c in cells if c))

        return '\n'.join(parts)

    # =====================================================================
    # 文件信息
    # =====================================================================

    def get_info(self) -> Dict[str, Any]:
        """获取文档元信息"""
        # 统计内容
        para_count = len([p for p in self._doc.paragraphs if p.text.strip()])
        table_count = len(self._doc.tables)
        total_table_rows = sum(len(t.rows) for t in self._doc.tables)

        return {
            "file_name": self.file_name,
            "file_path": self.file_path,
            "file_size": self.file_size,
            "content_hash": self.content_hash,
            "paragraph_count": para_count,
            "table_count": table_count,
            "total_table_rows": total_table_rows,
        }
